import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_docx(filename):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Colors
    PRIMARY_PINK = RGBColor(255, 45, 117)
    DARK_TEXT = RGBColor(20, 20, 30)
    MUTED_TEXT = RGBColor(100, 100, 120)
    CODE_COLOR = RGBColor(180, 20, 80)
    
    # Header Title
    title = doc.add_paragraph()
    title_run = title.add_run("DEZIR CLAB")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_PINK
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Complete Deployment Guide: Ninzahost (cPanel / VPS) & Vercel Free Hosting")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(13)
    sub_run.font.bold = True
    sub_run.font.color.rgb = DARK_TEXT
    sub.paragraph_format.space_after = Pt(16)

    # Callout Box: Summary
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "FFF0F5")
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    p = cell.paragraphs[0]
    r = p.add_run("📌 HOSTING COMPATIBILITY OVERVIEW\n")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = PRIMARY_PINK
    r2 = p.add_run(
        "NO, it is NOT compulsory to use Vercel! You can host Dezir Clab directly on your Ninzahost cPanel / Node.js hosting, VPS, or any server supporting Node.js. This guide provides exact instructions for both Ninzahost and Vercel."
    )
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def add_h1(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_PINK
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = DARK_TEXT
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            bp = p.add_run(bold_prefix)
            bp.font.bold = True
            bp.font.size = Pt(10)
            bp.font.color.rgb = DARK_TEXT
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_TEXT
        return p

    def add_code_block(code_text):
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0, 0)
        set_cell_background(c, "F4F4F8")
        set_cell_margins(c, top=120, bottom=120, left=160, right=160)
        p = c.paragraphs[0]
        r = p.add_run(code_text)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = CODE_COLOR
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. URL Architecture
    add_h1("1. How Platform URLs Work")
    add_body(
        "Dezir Clab is structured into distinct, isolated public, member, and secret administrative sections:",
        bold_prefix="Core URL Routing Structure: "
    )

    t = doc.add_table(rows=5, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Route Path", "Target Audience", "Functionality"]
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        set_cell_background(cell, "FF2D75")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("https://yourdomain.com/", "Public Visitors", "Sales page, brand features, invitation code entry button"),
        ("https://yourdomain.com/join/[token]", "Invited Members", "Account registration & token verification page"),
        ("https://yourdomain.com/dashboard", "Logged-in Members", "Private feed, creator stories rail, reactions & posts"),
        ("https://yourdomain.com/likecrazy", "Website Owner (Admin)", "Secret admin console (all old /admin links return 404)"),
    ]
    for row_idx, row_data in enumerate(data, start=1):
        bg = "FFFFFF" if row_idx % 2 == 1 else "F9F8FB"
        for col_idx, text in enumerate(row_data):
            cell = t.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.font.color.rgb = DARK_TEXT
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = PRIMARY_PINK

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2. Deploying on Ninzahost
    add_h1("2. Deploying on Ninzahost (cPanel / Node.js App)")
    add_body("Ninzahost provides cPanel with 'Setup Node.js App' (CloudLinux / Phusion Passenger). Here is the exact step-by-step setup:")

    add_h2("Step 1: Create Node.js Application in Ninzahost cPanel")
    add_body("1. Log in to your Ninzahost cPanel dashboard.")
    add_body("2. Scroll to the 'Software' section and click on 'Setup Node.js App'.")
    add_body("3. Click 'Create Application'.")
    add_body("4. Fill in the fields as follows:")
    add_body("   • Node.js version: Select 18.x or 20.x (Recommended: 20.x).")
    add_body("   • Application mode: Production.")
    add_body("   • Application root: dezir-clab (or community).")
    add_body("   • Application URL: Select your domain (e.g. yourdomain.com).")
    add_body("   • Application startup file: server.js (or run-server.js).")
    add_body("5. Click 'Create' in the top right corner.")

    add_h2("Step 2: Build & Upload the Code")
    add_body("1. On your local machine, run the standalone build command:")
    add_code_block("npm run build")
    add_body("2. This generates the .next/standalone folder containing a production-ready server.js and precompiled bundle.")
    add_body("3. Create a ZIP archive of your project (or upload via Git / FTP) containing:")
    add_body("   • .next/ (including .next/static and .next/standalone)")
    add_body("   • public/")
    add_body("   • package.json")
    add_body("   • prisma/ (with schema.prisma and dev.db or cloud postgres URL)")
    add_body("   • .env")
    add_body("4. In cPanel -> File Manager -> go into your application folder (e.g. /home/username/dezir-clab) and extract the files.")

    add_h2("Step 3: Copy Standalone Assets in cPanel")
    add_body("For standalone Next.js builds on cPanel, copy static assets into the standalone folder:")
    add_code_block("cp -r public .next/standalone/\ncp -r .next/static .next/standalone/.next/")

    add_h2("Step 4: Configure Environment Variables in Ninzahost cPanel")
    add_body("In 'Setup Node.js App' -> scroll to 'Environment variables' and add:")
    add_body("• DATABASE_URL: Your database connection string (SQLite file path or Neon/Supabase cloud PostgreSQL URI).")
    add_body("• NEXT_PUBLIC_APP_URL: https://yourdomain.com")
    add_body("• SESSION_SECRET: dezir_clab_secret_session_2026_production")
    add_body("• PORT: 3000 (or leave default for Passenger)")

    add_h2("Step 5: Run npm install & Start")
    add_body("1. In 'Setup Node.js App', click 'Run NPM Install'.")
    add_body("2. Click 'Restart Application'.")
    add_body("3. Visit https://yourdomain.com to see your website live!")

    # 3. Free Database Setup
    add_h1("3. Setting Up a 100% Free Cloud Database")
    add_body("You can either keep using SQLite on Ninzahost (zero setup required) or connect a free cloud PostgreSQL database (Neon.tech or Supabase).")
    
    add_h2("Option A: Free Cloud PostgreSQL on Neon.tech (Recommended)")
    add_body("1. Go to https://neon.tech and click 'Sign Up' (free tier: 0.5 GB storage, serverless autoscaling).")
    add_body("2. Create a project named 'dezir-clab' and copy the connection string:")
    add_code_block("postgresql://neondb_owner:YOUR_PASSWORD@ep-xyz-pooler.us-east-2.aws.neon.tech/neondb?sslmode=require")
    add_body("3. In prisma/schema.prisma, change provider to 'postgresql':")
    add_code_block('datasource db {\n  provider = "postgresql"\n  url      = env("DATABASE_URL")\n}')
    add_body("4. Run npx prisma db push && npm run seed to initialize your cloud DB.")

    add_h2("Option B: SQLite on Ninzahost (Zero External Database Needed)")
    add_body("You can also keep SQLite inside your Ninzahost folder. The database file is located at prisma/dev.db and stores all data right on your Ninzahost disk.")

    # 4. Hosting on Vercel
    add_h1("4. Alternative: 1-Click Free Hosting on Vercel")
    add_body("If you ever want instant free global CDN hosting with zero server maintenance:")
    add_body("1. Push your repository to GitHub (git push origin main).")
    add_body("2. Go to https://vercel.com and import your repository.")
    add_body("3. Add DATABASE_URL, NEXT_PUBLIC_APP_URL, and SESSION_SECRET in Environment Variables.")
    add_body("4. Click Deploy. Live in 60 seconds!")

    # 5. Admin Credentials & Operations
    add_h1("5. Admin Credentials & Daily Operations")
    add_body("Default Super Admin Login:", bold_prefix="Admin Access: ")
    add_code_block("URL: https://yourdomain.com/likecrazy\nEmail: admin@community.vip\nPassword: AdminSecret2026!")

    add_h2("Daily Administrator Checklist:")
    add_body("• Generating Invitations: Go to /likecrazy/invitations to generate private invite codes with custom usage limits (e.g. 1 use for VIP paying members).")
    add_body("• Posting Content: Go to /likecrazy/posts to publish photos, videos, and dispatches under any of the 10 resident personas with drag-and-drop upload.")
    add_body("• Messaging Members: Go to /likecrazy/messages to reply to member direct messages on behalf of any persona.")
    add_body("• Member Management: Go to /likecrazy/members to view, suspend, or ban user accounts.")

    # Save
    doc.save(filename)
    print(f"Successfully generated: {filename}")

if __name__ == "__main__":
    generate_docx("/Users/amit/community/Dezir_Clab_Deployment_and_Hosting_Guide.docx")
